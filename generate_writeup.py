"""
Script to generate the elaborated Carbon Accounting System write-up.
Run with: .venv/Scripts/python.exe generate_writeup.py
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document()


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')


def numbered(text):
    return doc.add_paragraph(text, style='List Number')


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


# ── TITLE ─────────────────────────────────────────────────────────────────────
t = doc.add_heading('Carbon Accounting System', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── OVERVIEW ──────────────────────────────────────────────────────────────────
h1('Overview')

body(
    'This CAS serves as an accounting ledger for CO₂ at Octavia Carbon. The ledger accounts for the '
    'CO₂ that is captured and stored versus the CO₂ emissions in the process of capture and storage. '
    'The system is updated weekly using aggregated cycle data and determines whether the facility is '
    'net carbon positive or net carbon negative at any given point.'
)

body(
    'The capture process accounts for CO₂ at the last capture monitoring point, which is the point '
    'at which the CO₂ is liquefied. Emissions are divided into two categories: embodied emissions '
    'and operational emissions. The operational emissions are accounted for up to the same last '
    'monitoring point.'
)

# ── CO2 CAPTURED ──────────────────────────────────────────────────────────────
h1('CO₂ Captured')

body(
    'The system monitors CO₂ at four distinct points along the capture process. These points allow '
    'the system to track CO₂ mass at each stage transition and calculate inter-stage losses. The '
    'four monitoring points are:'
)

numbered('CO₂ Adsorbed (ADS)')
numbered('CO₂ Desorbed (DES)')
numbered('CO₂ Collected in the Balloon Bag (BAG)')
numbered('CO₂ Liquefied (LIQ)')

body(
    'These observation points serve two purposes: they feed the efficiency monitoring chain and '
    'they provide the final captured CO₂ figure at the last monitoring point (LIQ) which is what '
    'enters the carbon accounting equation.'
)

h2('CO₂ Adsorbed (ADS)')

body(
    'The ADS figure is determined by integrating the difference between inlet and outlet CO₂ '
    'concentrations over the duration of the adsorption cycle. CO₂ sensors measure concentration '
    'continuously at both the inlet and outlet of each module. As the sorbent saturates, the '
    'outlet concentration rises back toward inlet levels — this is the breakthrough point. '
    'The mass of CO₂ adsorbed per cycle is calculated as:'
)

equation('CO₂ Adsorbed = ∫ (C_in − C_out) × Q  dt')

body(
    'Where C_in is the inlet CO₂ concentration, C_out is the outlet CO₂ concentration, and Q is '
    'the volumetric air flow rate. The integral is taken over the full duration of the adsorption '
    'cycle. This is calculated per cycle and logged by SCADA.'
)

h2('CO₂ Desorbed (DES)')

body(
    'Once the adsorption phase ends, steam is applied to the sorbent bed to release the bound CO₂. '
    'The DES figure is determined by measuring the corrected CO₂ gas flow from the reactor during '
    'desorption and integrating it over the desorption period using the Ideal Gas Law to convert '
    'from volumetric to mass units:'
)

equation('n = PV / RT     →     m = n × M_CO₂  (44.01 g/mol)')

body(
    'The ratio of DES to ADS defines the ADS→DES efficiency, which captures losses at the desorption '
    'stage such as incomplete release or measurement error.'
)

h2('CO₂ Collected in the Balloon Bag (BAG)')

body(
    'The desorbed CO₂ flows into a balloon bag — a flexible buffer vessel between the reactor '
    'desorption cycle and the downstream liquefaction unit. The CO₂ measured at the bag quantifies '
    'the separation losses between the reactor outlet and the bag inlet. The ratio of BAG to DES '
    'defines the DES→BAG efficiency, capturing any gas-handling or separation losses in that section.'
)

h2('CO₂ Liquefied (LIQ) — The Last Monitoring Point')

body(
    'The cryogenic storage tank is fitted with load cells that continuously measure the total weight '
    'of liquefied CO₂. The increase in weight over a cycle or a week directly gives the mass of CO₂ '
    'that has been successfully captured and stored. This is the figure that enters the carbon '
    'accounting equation as CO₂ Captured.'
)

body(
    'Using the last monitoring point as the captured CO₂ figure means the system only claims credit '
    'for CO₂ that has passed through the full process chain. All upstream losses are automatically '
    'excluded. The ratio of LIQ to BAG defines the BAG→LIQ efficiency.'
)

body(
    'Where the liquefied CO₂ figure is not yet available for a given week — for example if the '
    'operator has not yet entered the load cell reading — the system falls back to using the '
    'BAG figure as the gross captured CO₂, and stage 3 loss is recorded as zero until the '
    'liquefaction data is entered.'
)

h2('Efficiency Chain')

body(
    'The four monitoring points create a three-stage efficiency chain that spans the full process:'
)

bullet('ADS → DES: adsorption-to-desorption efficiency — losses at the desorption stage')
bullet('DES → BAG: desorption-to-bag efficiency — losses in gas handling and separation')
bullet('BAG → LIQ: bag-to-liquefaction efficiency — losses in compression and liquefaction')
bullet('ADS → LIQ: overall system efficiency — total losses from first capture to final storage')

body(
    'Each stage efficiency is computed as a ratio of the downstream CO₂ figure to the upstream '
    'CO₂ figure for that stage. These are tracked weekly and displayed on the dashboard as a '
    'CO₂ flow waterfall. The total system loss is:'
)

equation('Total Loss = ADS − Gross Captured (LIQ or BAG)')

body(
    'Capture data is aggregated weekly by summing all cycle records within the ISO week period.'
)

# ── CO2 EMISSIONS ─────────────────────────────────────────────────────────────
h1('CO₂ Emissions')

body(
    'The system calculates two categories of CO₂ emissions: operational emissions and embodied '
    'emissions. The total is deducted from the gross CO₂ captured to produce the net figure.'
)

equation('CO₂ EMISSIONS = OPERATIONAL EMISSIONS + EMBODIED EMISSIONS')

h2('Operational Emissions')

body(
    'Operational emissions are the day-to-day emissions from running the capture process. They scale '
    'with plant activity and are aggregated weekly from SCADA cycle data. There are three components:'
)

h3('Thermal Emissions (Boiler)')

body(
    'The boiler generates steam for sorbent desorption and is electrically powered. Boiler energy '
    'consumption is logged per cycle by SCADA (boiler_kwh). The thermal emissions are calculated as:'
)

equation('Thermal Emissions = Boiler kWh × Grid EF')

body(
    'The grid emission factor is set at 0.049 kg CO₂/kWh (Kenya Power 2024). This value is '
    'stored in the system_config table and is configurable by an administrator.'
)

h3('Auxiliary Emissions (Non-Thermal Electrical)')

body(
    'Auxiliary emissions cover all other electrical consumption in the process. This includes the '
    'fans/blowers (NM1–NM4), the SRV/LRVP (Steam Recovery Valve / Liquid Ring Vacuum Pump), and '
    'the cooling tower. Each is logged separately per cycle by SCADA. Additionally, the energy '
    'consumed by the liquefaction unit and CO₂ transfer operations — obtained from the weekly '
    'SCADA operating summary and entered manually — is added to the auxiliary total.'
)

equation('Auxiliary Emissions = (SRV/LRVP + CT + Fans + Liquefaction kWh) × Grid EF')

body(
    'If SCADA provides a total kWh figure but no component breakdown, the system estimates the '
    'split using default ratios: 70% thermal (boiler) and 30% auxiliary. The total operational '
    'emissions are then:'
)

equation('Total Operational Emissions = (Boiler kWh + Auxiliary kWh) × Grid EF')

equation('Total Operational Emissions = Total kWh × Grid EF')

h3('Sorbent Degradation Emissions')

body(
    'The sorbent material requires periodic re-functionalisation using PEI and methanol. The '
    'upstream lifecycle emissions of these materials are accounted for as an operational emission '
    'rate applied proportionally to the CO₂ captured each week. The rates, sourced from the '
    'Carbon Nest LCA Workbook (IPCC 2021 GWP100), are:'
)

bullet('PEI: 0.0688 kg CO₂eq per kg CO₂ captured')
bullet('Methanol: 0.0834 kg CO₂eq per kg CO₂ captured')
bullet('Combined rate: 0.1522 kg CO₂eq per kg CO₂ captured')

equation('Sorbent Emissions = Gross Captured (kg) × 0.1522')

body(
    'Note: sorbent degradation emissions are currently tracked as a rate in the system configuration '
    'but are not yet broken out as a separate line in the weekly summary table — they are factored '
    'into the overall operational emissions total via the LCA-derived rate.'
)

h2('Embodied Emissions')

body(
    'Embodied emissions represent the one-time CO₂ costs of constructing and assembling the '
    'facility, including civil works, structural steel, aluminium, mild steel, shipping containers, '
    'concrete, mechanical equipment, instrumentation, the liquefaction system, and the baseline '
    'sequestration foregone from the land on which the facility sits. These are calculated once '
    'via life cycle assessment and then amortised over the plant lifetime.'
)

body(
    'The total lifetime embodied emissions figure is 67,267.5 kg CO₂eq, sourced from the '
    'Carbon Nest LCA Workbook. The plant has a design lifetime of 10 years (520 weeks) and '
    'a design lifetime capture of 250,000 kg CO₂ (25 t/yr × 10 yr).'
)

h3('Amortisation — Ramp-Up Phase')

body(
    'The plant is currently in an early ramp-up phase where it is not yet operating at full design '
    'capacity. The amortisation model allocates only 26% of the total lifetime embodied emissions '
    'to this ramp-up phase, spread over a 5-year (260-week) period. This prevents the full '
    'embodied emissions burden from being charged against a period of lower-than-design output. '
    'The calculation is as follows:'
)

equation('Ramp-up embodied total = 67,267.5 × 0.26 = 17,489.6 kg CO₂eq')
equation('Weekly embodied (ramp-up) = 17,489.6 ÷ 260 weeks = 67.27 kg CO₂eq/week')

body(
    'For reference, if the full lifetime embodied were spread evenly across all 520 weeks, the '
    'weekly charge would be 67,267.5 ÷ 520 = 129.36 kg CO₂eq/week. The ramp-up model therefore '
    'charges approximately half of what the steady-state rate would be.'
)

body(
    'The weekly embodied figure is constant — it does not vary with output, number of cycles, '
    'or any operational variable. It is applied as a fixed deduction every week regardless of '
    'how much CO₂ was captured.'
)

equation('Weekly Embodied Emissions = 67.27 kg CO₂eq/week  (ramp-up phase, current)')

# ── CARBON ACCOUNTING EQUATION ────────────────────────────────────────────────
h1('Carbon Accounting Equation')

body(
    'The system is updated weekly using the aggregated cycle data to calculate the net CO₂ captured. '
    'The core equation is:'
)

equation('NET CO₂ = CO₂ CAPTURED – CO₂ EMISSIONS')
equation('CO₂ EMISSIONS = EMBODIED EMISSIONS + OPERATIONAL EMISSIONS')

body(
    'CO₂ Captured is the liquefied CO₂ figure from the load cells (or the BAG figure if liquefaction '
    'data is not yet available). CO₂ Emissions is the sum of the weekly embodied charge (67.27 kg) '
    'plus the total operational emissions calculated from the week\'s energy and sorbent data. '
    'Expanded:'
)

equation('NET CO₂ = LIQ – [(Total kWh × Grid EF) + Sorbent Emissions + 67.27]')

body(
    'A positive NET CO₂ figure means the facility removed more CO₂ from the atmosphere than it '
    'emitted that week — it is net carbon positive and carbon removal credits can be issued for '
    'that quantity. A negative NET CO₂ figure means emissions exceeded capture — the facility '
    'is a net emitter for that period and no credits are issued.'
)

body(
    'The system also tracks a cumulative NET CO₂ figure across all weeks of operation, which '
    'is the running total of net CO₂ removed since the plant started. Both the weekly and '
    'cumulative figures are displayed on the dashboard.'
)

# ── DATA LOGGING ──────────────────────────────────────────────────────────────
h1('Data Logging')

body(
    'Data is sourced from SCADA and fed to the system in CSV format. There are two primary '
    'CSV inputs: the plant cycle CSV and the energy CSV. The operator selects the week the data '
    'is for and the system automatically aggregates cycle records to weekly totals. Manual entry '
    'is also available for any data point.'
)

h2('Plant Cycle CSV')

body(
    'The cycle CSV contains per-cycle operational data. Each row is one complete adsorption-desorption '
    'cycle for one module pair. The key columns are:'
)

bullet('Cycle # — cycle identifier')
bullet('Machine — module pair identifier (e.g., 1n3, 2n4, NM1, NM3)')
bullet('ADS CO₂ (kg) — CO₂ adsorbed from the breakthrough curve integration')
bullet('DES CO₂ (kg) — CO₂ desorbed, calculated via corrected gas flow and ideal gas law')
bullet('BAG CO₂ (kg) — CO₂ collected in the balloon bag')
bullet('eTotal kWh — total electrical energy for the cycle')
bullet('Steam (kg) — steam consumed in the desorption phase')

body(
    'The cycle CSV does not include the liquefied CO₂ data. This is obtained separately from '
    'the weekly SCADA operating summary and entered manually by the operator alongside the '
    'liquefaction energy.'
)

h2('Energy CSV')

body(
    'The energy CSV provides a granular breakdown of electrical consumption by equipment category, '
    'complementing the total kWh figure in the cycle CSV. The key columns are:'
)

bullet('Boiler kWh — electrical energy consumed by the steam boiler')
bullet('SRV/LRVP kWh — Steam Recovery Valve / Liquid Ring Vacuum Pump energy')
bullet('CT kWh — cooling tower energy')
bullet('Fan kWh (NM1–NM4) — individual fan/blower energy per module')

body(
    'The system merges cycle and energy data on import. If a total kWh figure is present but the '
    'component breakdown is missing or zero, the system estimates the split at 70% boiler / 30% '
    'auxiliary. Duplicate cycles are detected and rejected on import to prevent double-counting.'
)

h2('Manual Entry and the Liquefied CO₂ Figure')

body(
    'The operator can choose to input any data manually through the data entry interface if the '
    'CSV option is unavailable or misbehaves. The liquefied CO₂ figure from the load cells is '
    'always entered manually, as it is not exported in the standard SCADA cycle CSV — it is '
    'read from the weekly SCADA summary report. The operator also enters the liquefaction '
    'energy (kWh) at the same time. These two values are stored against the selected ISO week '
    'and are used when calculating the weekly summary.'
)

h2('Database')

body(
    'All data is stored in a PostgreSQL database. The core tables are:'
)

bullet('cycle_data — raw per-cycle records (CO₂ at ADS/DES/BAG, energy by component, steam, module pair)')
bullet('weekly_summary — aggregated weekly calculations (net removal, total emissions, losses, energy intensity)')
bullet('embodied_infrastructure — LCA component data for infrastructure embodied emissions')
bullet('system_config — configurable parameters (grid emission factor, geothermal EF, etc.)')
bullet('audit_log — full audit trail of all data changes with timestamp and user')

body(
    'Weekly aggregation is performed by summing all cycle_data records where the cycle start_time '
    'falls within the ISO week\'s Monday-to-Sunday window. The weekly embodied figure is then '
    'applied as a fixed addition and the calculate_weekly_metrics function computes all '
    'derived values from the aggregated inputs.'
)

# ── DATA VISUALISATION AND REPORTS ────────────────────────────────────────────
h1('Data Visualisation and Reports')

body(
    'The system has a dynamic dashboard. The data is represented weekly and the dashboard is '
    'dynamic to a weekly level — the user selects the week they want and all charts update '
    'accordingly. The dashboard can also be filtered by module pair (1n3 or 2n4) to isolate '
    'performance for each pair independently.'
)

h2('Dashboard')

body(
    'The dashboard contains the following components:'
)

bullet(
    'Weekly Performance Summary Cards — NET CO₂, CO₂ CAPTURED (LIQUEFIED), TOTAL EMISSIONS, '
    'and TOTAL LOSSES (ADS→LIQ), displayed as headline KPI cards.'
)
bullet(
    'CO₂ Flow Analysis — shows the CO₂ mass at each of the four monitoring points from ADS to LIQ, '
    'with inter-stage losses shown as both absolute kg values and efficiency percentages. '
    'Visualised as a waterfall/flow chart.'
)
bullet(
    'Carbon Balance Waterfall — steps through the carbon accounting equation from gross captured '
    'down through each emission deduction (embodied, thermal, auxiliary) to arrive at the net figure, '
    'making the contribution of each category immediately visible.'
)
bullet(
    'Energy Breakdown — weekly electrical consumption disaggregated by equipment category '
    '(boiler, fans, SRV/LRVP, cooling tower, liquefaction).'
)
bullet(
    'Energy Intensity — total energy consumed (kWh) per tonne of CO₂ captured, plotted as a '
    'trend over time. Calculated as: Energy Intensity = Total kWh ÷ (Gross Captured kg ÷ 1000).'
)
bullet(
    'CO₂ Losses Chart — visualises the losses at each stage of the process to identify where '
    'the most significant losses are occurring.'
)

h2('Reports')

body(
    'Reports can be downloaded in CSV and Excel formats from the reports page. The available '
    'report types are:'
)

bullet(
    'Weekly Summaries — one row per week, containing CO₂ captured, each emission category, '
    'net removal, energy totals, energy intensity, and loss figures.'
)
bullet(
    'Cycle Data — full granular per-cycle records for the selected date range.'
)
bullet(
    'Module Pair Comparison — a side-by-side breakdown of performance for Module 1&3 versus '
    'Module 2&4. Includes per-pair CO₂ totals, stage efficiencies, energy per kg CO₂, average '
    'per-cycle figures, and the overall efficiency ratio between the two pairs. '
    'The pair contribution to total CO₂ captured is also reported as a percentage.'
)

# ── OPTIMISATION ──────────────────────────────────────────────────────────────
h1('Optimisation')

body(
    'The system has a simulations page with three tools: a geothermal scenario calculator, a '
    'historical what-if analysis, and a Monte Carlo simulation. These tools allow the team to '
    'model the carbon impact of operational changes and quantify uncertainty in the accounting figures.'
)

h2('Geothermal Energy Scenario')

body(
    'The geothermal scenario models what happens to the carbon balance when the steam boiler — '
    'currently electrically powered from the grid — is replaced by a direct geothermal steam supply. '
    'The key change in the calculation is that the thermal emissions term is no longer '
    'Boiler kWh × Grid EF but instead Steam kg × Geothermal Steam EF:'
)

equation('Current Thermal Emissions  = Boiler kWh × Grid EF  (0.049 kg CO₂/kWh)')
equation('Geothermal Thermal Emissions = Steam kg × Geothermal Steam EF  (0.005 kg CO₂/kg steam)')

body(
    'Auxiliary emissions remain on the grid in both scenarios — only the thermal/steam component '
    'changes. The geothermal steam emission factor of 0.005 kg CO₂/kg steam is the default, '
    'reflecting the low fugitive non-condensable gas emissions from geothermal wells, and is '
    'configurable in the admin panel. The full scenario calculation is:'
)

equation('Geo Operational = (Steam kg × Geo Steam EF) + (Auxiliary kWh × Grid EF)')
equation('Geo Net = CO₂ Captured – (Geo Operational + Embodied)')
equation('Improvement = Geo Net – Current Net')

body(
    'The historical what-if tab applies this recalculation to every week in the database and '
    'shows how many weeks would have been net positive under geothermal versus the current '
    'grid scenario, along with the total net improvement across all tracked weeks.'
)

h2('Monte Carlo Uncertainty Analysis')

body(
    'The Monte Carlo simulation runs N iterations of the carbon accounting calculation, each time '
    'sampling the key input parameters from uniform distributions within user-defined min/max '
    'ranges. The sampled parameters are:'
)

bullet('Boiler energy consumption (kWh) — min/max range')
bullet('Auxiliary electrical energy (kWh) — min/max range')
bullet('Steam consumed (kg) — min/max range for geothermal scenario')
bullet('CO₂ collected (kg) — min/max range')

body(
    'For each iteration, the full geothermal scenario calculation is run using the sampled values, '
    'producing a distribution of net removal outcomes for both the current and geothermal scenarios. '
    'The outputs reported are:'
)

bullet('Mean net removal (current and geothermal)')
bullet('Standard deviation of net removal')
bullet('Percentage of simulations where the plant is net positive')
bullet('Min/max/mean net improvement from geothermal substitution')

body(
    'Results are visualised as overlapping histograms and box plots for both scenarios, with a '
    'break-even line at zero. This provides the probability bounds needed for defensible carbon '
    'credit claims under verification standards.'
)

# ── SYSTEM AND ACCESS ─────────────────────────────────────────────────────────
h1('System and Access')

body(
    'The CAS is a multi-page Streamlit web application backed by a PostgreSQL database. It is '
    'accessible via browser and requires authentication. Two user roles are defined: User (view-only '
    'access to dashboard, reports, and simulations) and Administrator (full access including data '
    'import, system configuration, and user management).'
)

body(
    'System configuration parameters — including the grid emission factor, geothermal steam EF, '
    'and other constants — are stored in the system_config table and are editable by an administrator '
    'through the admin panel. All changes are recorded in the audit log with a timestamp and the '
    'username of the user who made the change.'
)

# ── SAVE ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\user\Desktop\CASS\Carbon Accounting System Write Up.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
